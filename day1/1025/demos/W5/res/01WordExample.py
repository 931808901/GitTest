import docx #python-docx
import win32com
import win32com.client
from docx.shared import Inches
from openpyxl.reader.excel import load_workbook


def writeWordDoc():
    # 调用win32com的系统接口打开Word
    word = win32com.client.Dispatch("Word.Application")  # 操作word
    word.Visible = False  # 可以看见

    # 构造一个新的文档对象
    doc = word.Documents.Add()  # 插入文档

    # 从头部插入5行文本
    rng = doc.Range(0, 0)  # 操作位置，从00开始
    rng.InsertAfter("锄禾日当午\n")
    rng.InsertAfter("汗滴禾下午\n")
    rng.InsertAfter("一本小破书\n")
    rng.InsertAfter("一看一下午\n")
    rng.InsertAfter("  ——悯码农")

    # 写入指定路径的文件中
    filename = "C:\\Users\\idea\\Desktop\\" + "悯码农" + ".doc"
    doc.SaveAs(filename)  # 保存
    doc.Close(True)  # Flase强行关闭,True等待缓冲区写入完毕

    # 退出Word
    word.Application.Quit()  # 退出系统接口

def readWordDoc():
    doc = docx.Document("C:\\Users\\idea\\Desktop\\悯码农.doc")
    # doc = docx.Document("demo.docx")
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    print("\n".join(fullText))


if __name__ == '__main__':
    # writeWordDoc()
    readWordDoc()
    pass
