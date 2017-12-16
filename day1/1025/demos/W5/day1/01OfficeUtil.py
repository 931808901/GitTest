#! C:\Python36\python.exe
# coding:utf-8
'''
about what
'''
import docx
import win32com.client
from openpyxl import load_workbook


def writeWord(text, filepath):
    # 调用win32com的系统接口打开Word
    word = win32com.client.Dispatch("Word.Application")  # 操作word
    word.Visible = False  # 可以看见

    # 构造一个新的文档对象
    doc = word.Documents.Add()  # 插入文档

    # 从头部插入5行文本
    rng = doc.Range(0, 0)  # 操作位置，从00开始
    rng.InsertAfter(text)

    # 写入指定路径的文件中
    doc.SaveAs(filepath)  # 保存
    doc.Close(True)  # Flase强行关闭,True等待缓冲区写入完毕

    # 退出Word
    word.Application.Quit()  # 退出系统接口

def readWord(filepath):
    doc = docx.Document(filepath)
    # doc = docx.Document("demo.docx")
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText

def writeExcel(datalist, filepath):
    # 调用系统接口打开Excel
    excel = win32com.client.Dispatch("Excel.Application")
    excel.Visible = False  # 程序窗口不可见

    # 新建工作表并定位到当前sheet页
    wb = excel.Workbooks.Add()
    currentSheet = wb.ActiveSheet

    #
    rows = len(datalist)
    keylist = list(datalist[0].keys())
    cols = len(keylist)

    # 将keylist中的键名写入第一行
    for i in range(1, len(keylist) + 1):
        currentSheet.Cells(1, i).value = keylist[i - 1]

    # 对指定行列的单元格赋值
    for row in range(2, rows + 2):
        for col in range(1, cols + 1):
            currentSheet.Cells(row, col).value = datalist[row - 2][keylist[col - 1]]

    # 保存内存数据到文件，关闭IO流
    wb.SaveAs(filepath)
    wb.Close(True)

    # 关闭Excel
    excel.Application.Quit()

def readExcel(filepath):
    # 载入指定路径的Excel文件到内存
    wb = load_workbook(filepath)

    # 拿到文件中的所有工作表
    sheetNames = wb.get_sheet_names()
    # 打印第一个工作表的信息
    sheet = wb.get_sheet_by_name(sheetNames[0])

    datalist = []
    keylist = []

    # 根据行列索引遍历每一个单元格内容
    for row in range(1, sheet.max_row + 1):
        if row != 1:
            mdict = {}
        for col in range(1, sheet.max_column + 1):
            val = sheet.cell(row=row, column=col).value

            if row == 1:
                keylist.append(val)
            else:
                mdict[keylist[col-1]] = val

        if row > 1:
            datalist.append(mdict)

    return datalist

if __name__ == "__main__":
    # writeWord("我牛不牛逼我不知道\n但是当别人跟我说"你死了地球照样转的时候"\n我听着就觉得地球在硬撑；",r"C:\Users\idea\Desktop\什么是自信.doc")

    # textlist = readWord(r"C:\Users\idea\Desktop\什么是自信.doc")
    # print(textlist)

    # datalist = [
    #     {"name": "张三", "age": 20, "hobby": "看片"},
    #     {"name": "lisi", "age": 20, "hobby": "coding"},
    #     {"name": "wangwu", "age": 40, "hobby": "读书"},
    #     {"name": "zhaoliu", "age": 20, "hobby": "coding"},
    # ]
    # writeExcel(datalist,r"C:\Users\idea\Desktop\人员信息.xlsx")

    print(readExcel(r"C:\Users\idea\Desktop\人员信息.xlsx"))

    print("main over")
